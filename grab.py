#!/usr/bin/python

import requests
import re
from docx import Document

response = requests.get('http://xxx.lanl.gov/list/AI.ai/pastweek?skip=0&show=25')
str = response.text
url = re.findall(r"href=\"(/list/AI.ai/pastweek\?show=[0-9]+)\"", str)
responseall = requests.get("http://xxx.lanl.gov" + url[0])
strall = responseall.text
urlall = re.findall(r"href=\"(/abs/[0-9]+.[0-9]+)\"", strall) 

document = Document()
count = 0
for item in urlall:
    count = count + 1
    print item
    abstract = requests.get('http://xxx.lanl.gov' + item)
    text = abstract.text
    name = re.findall(r"<a href=\"http://arxiv.org/find/AI.*\">(.*)</a>",text)
    title = re.findall(r"<span class=\"descriptor\">Title:</span>\n(.*)</h1>", text)
    content = re.findall(r"<span class=\"descriptor\">Abstract:</span>(.*)\n</blockquote>", text, re.S)
    subject = re.findall(r"<span class=\"primary-subject\">(.*)</span>(.*)</td>" , text)
    document.add_heading(title, 1)

    temp = ",".join(name)
    #print temp
    para1 = document.add_paragraph(temp)
    para = document.add_paragraph(content)
    #print subject
    temp1 = "".join(v for v in subject[0])
    #print temp1
    para2 = document.add_paragraph(temp1)
    para3 = document.add_paragraph('http://xxx.lanl.gov' + item)
    document.save('a.docx')
print count
